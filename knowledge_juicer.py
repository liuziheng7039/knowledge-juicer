import streamlit as st
import fitz  # PyMuPDF - 用来读取 PDF 文件
import dashscope
from dashscope.api_entities.dashscope_response import Role
import base64
from docx import Document
import io
import random
import json
import re
import threading
import time
import os

# =========================================================
# 应用说明
# =========================================================
# 这是一个 AI 学习助手，功能包括：
# 1. 上传课件 PDF，自动提取核心知识点
# 2. 生成结构化的知识清单（可下载 Word 文档）
# 3. 根据知识点自动出题（单选+简答+论述）
# 4. 答错的题会换个问法重新考，直到掌握
# 5. 基于知识清单的智能答疑
#
# 主要特点：
# - 0基础模式：知识点多、解释详细、题目简单
# - 高玩模式：知识点精炼、题目难度高
# - 上传真题后会自动标记高频考点

# =========================================================
# 界面文案配置
# =========================================================
UI = {
    # 页面标题
    "PAGE_TITLE": "榨知机 V1.6",
    "SIDEBAR_TITLE": "🍋 榨知机 V1.6",
    "MAIN_TITLE": "🍋 榨知机 V1.6：你的期末救星",

    # 开发者模式提示
    "DEV_MODE": "✅ 开发者演示模式",

    # 工具箱按钮
    "TOOLBOX": "🛠️ 工具箱",
    "QUIZ_BTN": "🙋‍♂️ 考考我（来一组题练练手）",
    "DOWNLOAD_BTN": "📄 下载 知识清单",

    # 设置区域
    "SETTINGS": "### ⚙️ 复习设置",
    "GOAL": "目标：",
    "MODE_BEGINNER": "俺0基础 ！（召唤所有知识）",
    "MODE_PRO": "我是高玩！（直接上最难的）",
    "UPLOAD_MAIN": "上传课件 (必需)",
    "UPLOAD_EXAM": "上传真题 (可选，用于标记考点)",
    "START_BTN": "🚀 开始学习",

    # 生成过程提示
    "GEN_PROGRESS": "正在生成知识清单...",
    "GEN_ESTIMATE": "预计还需 {time} 分钟",
    "GEN_SUCCESS_TEMP": "🎉 知识清单生成成功！",
    "GEN_SUCCESS": "✅ 《{course_name}》知识清单已生成！侧边栏可下载/考考我。",
    "GEN_FAIL": "知识清单生成失败：",

    # 知识清单展示
    "LEFT_HEADER": "📄 核心知识清单",
    "MD_TITLE": "# 📘 核心知识清单",
    "MD_EXPLAIN": "- 解释：",
    "MD_EXAMPLE": "- 例子：",
    "MD_SCORE": "（重点指数：{score}）",

    # 右侧功能区
    "RIGHT_HEADER": "🤖 右侧：答疑 / 刷题与批改",
    "MODE_SWITCH": "模式",
    "MODE_QA": "答疑",
    "MODE_QUIZ": "刷题与批改",

    # 答疑功能
    "QA_INPUT": "基于知识清单提问…（不会回答知识清单未覆盖的新考点）",
    "QA_SPINNER": "AI 正在基于知识清单作答...",

    # 刷题功能
    "QUIZ_HINT": "点击侧边栏「考考我」开始刷题",
    "QUIZ_TITLE": "### 📝 高频卷（10题）",
    "QUIZ_PROGRESS": "进度：第 {cur}/{total} 题｜题组：{set_id}",
    "QUIZ_CONCEPT": "考点：{title}",
    "QUIZ_STEM": "**题目：**",

    # 错题循环
    "REMEDIAL_TITLE": "### 🔁 错题循环（直到掌握）",
    "REMEDIAL_META": "考点：{title} ｜ 已错次数：{n}/4",
    "REMEDIAL_STUCK": "该考点已连续错误达到上限（4次）。建议回看该概念后再挑战。",
    "REMEDIAL_REVIEW": "📌 建议回看：",
    "REMEDIAL_DONE": "🎉 错题已全部掌握！本轮结束。",
    "REMEDIAL_END": "🎉 错题循环已结束（部分考点需回看后再挑战）。",

    # 批改反馈
    "GRADE_FAIL": "⚠️ 批改失败：",
    "CORRECT": "✅ 正确",
    "WRONG": "❌ 错误",
    "RIGHT_ANSWER": "正确答案：",
    "SUGGEST_REVIEW": "建议回看该概念。",

    # 下一题按钮
    "NEXT_BTN": "➡️ 下一题 / 继续",

    # 简答题输入
    "TEXTAREA_LABEL": "请输入你的答案：",
    "TEXTAREA_PLACEHOLDER": "简答/论述：写要点即可，越结构化越好。",
    "SUBMIT_BTN": "✅ 提交答案",
    "NEED_INPUT": "请先输入答案。",

    # 未生成提示
    "EMPTY_HINT": "👈 请先在左侧上传课件并点击「开始学习」。生成知识清单后才会解锁：下载 / 高频刷题 / 答疑。",

    # 上传提示
    "UPLOAD_TIP": "把 PDF 拖到下面区域，或点击浏览按钮上传。",
}

# =========================================================
# 页面配置
# =========================================================
st.set_page_config(page_title=UI["PAGE_TITLE"], page_icon="🍋", layout="wide")

# 自定义 CSS 样式 - 全面美化界面
st.markdown("""
<style>
    /* ===== 全局背景 ===== */
    .stApp {
        background: linear-gradient(160deg, #f8f9fc 0%, #eef1f8 100%);
    }

    /* ===== 隐藏默认元素，更简洁 ===== */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* ===== 侧边栏 ===== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e1e2f 0%, #1a1a3e 50%, #12234d 100%) !important;
    }
    [data-testid="stSidebar"] [data-testid="stMarkdown"],
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stRadio label,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] .stCaption {
        color: #c8cfe0 !important;
    }
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {
        color: #ffffff !important;
    }
    [data-testid="stSidebar"] hr {
        border-color: rgba(255,255,255,0.1) !important;
    }
    [data-testid="stSidebar"] .stButton > button {
        background: linear-gradient(135deg, #6c63ff 0%, #8b5cf6 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.6rem 1.2rem !important;
        font-weight: 600 !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 14px rgba(108, 99, 255, 0.35) !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(108, 99, 255, 0.5) !important;
    }
    [data-testid="stSidebar"] .stDownloadButton > button {
        background: linear-gradient(135deg, #34d399 0%, #06b6d4 100%) !important;
        color: #fff !important;
        border: none !important;
        border-radius: 12px !important;
        font-weight: 600 !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 14px rgba(52, 211, 153, 0.35) !important;
    }
    [data-testid="stSidebar"] .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(52, 211, 153, 0.5) !important;
    }
    [data-testid="stSidebar"] [data-testid="stFileUploader"] {
        border: 2px dashed rgba(108, 99, 255, 0.4) !important;
        border-radius: 12px !important;
        background: rgba(255,255,255,0.03) !important;
        transition: border-color 0.3s ease !important;
    }
    [data-testid="stSidebar"] [data-testid="stFileUploader"]:hover {
        border-color: rgba(108, 99, 255, 0.7) !important;
    }

    /* ===== 进度条 ===== */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #6c63ff 0%, #8b5cf6 40%, #a78bfa 70%, #c4b5fd 100%) !important;
        border-radius: 10px;
    }

    /* ===== 主区域标题 ===== */
    .main-hero-title {
        font-size: 2.2rem;
        font-weight: 800;
        background: linear-gradient(135deg, #6c63ff, #8b5cf6, #a855f7);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 0.2rem;
        letter-spacing: -0.5px;
    }
    .main-hero-sub {
        font-size: 1rem;
        color: #8892a4;
        margin-bottom: 1.5rem;
    }

    /* ===== 成功提示动画 ===== */
    @keyframes fadeInScale {
        0% { opacity: 0; transform: scale(0.8); }
        100% { opacity: 1; transform: scale(1); }
    }
    .success-banner {
        animation: fadeInScale 0.5s ease-out;
    }

    /* ===== 知识卡片 ===== */
    .kc-card {
        background: #ffffff;
        border-radius: 14px;
        padding: 1.2rem 1.4rem;
        margin-bottom: 0.85rem;
        box-shadow: 0 1px 8px rgba(108, 99, 255, 0.06), 0 1px 3px rgba(0,0,0,0.04);
        border-left: 4px solid #6c63ff;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .kc-card:hover {
        transform: translateX(4px);
        box-shadow: 0 4px 20px rgba(108, 99, 255, 0.12), 0 2px 6px rgba(0,0,0,0.06);
    }
    .kc-card.star {
        border-left-color: #f472b6;
        background: linear-gradient(135deg, #ffffff 0%, #fdf2f8 100%);
    }
    .kc-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 0.5rem;
        gap: 0.8rem;
    }
    .kc-title {
        font-size: 1rem;
        font-weight: 700;
        color: #1e1e2f;
        margin: 0;
        flex: 1;
        line-height: 1.4;
    }
    .kc-badge {
        display: inline-flex;
        align-items: center;
        padding: 3px 10px;
        border-radius: 20px;
        font-size: 0.72rem;
        font-weight: 700;
        color: white;
        white-space: nowrap;
        flex-shrink: 0;
    }
    .kc-badge.hot { background: linear-gradient(135deg, #f472b6, #ef4444); }
    .kc-badge.warm { background: linear-gradient(135deg, #8b5cf6, #6c63ff); }
    .kc-badge.cool { background: linear-gradient(135deg, #60a5fa, #3b82f6); }
    .kc-explain {
        font-size: 0.88rem;
        color: #4b5563;
        line-height: 1.7;
        margin-bottom: 0.3rem;
    }
    .kc-explain .lbl {
        color: #6c63ff;
        font-weight: 600;
        font-size: 0.82rem;
    }
    .kc-example {
        background: #f5f3ff;
        border-radius: 8px;
        padding: 0.5rem 0.8rem;
        margin-top: 0.35rem;
        font-size: 0.84rem;
        color: #6b7280;
        line-height: 1.6;
        border: 1px solid #ede9fe;
    }

    /* ===== 欢迎页 ===== */
    .welcome-container {
        text-align: center;
        padding: 3rem 2rem;
        background: #ffffff;
        border-radius: 20px;
        box-shadow: 0 2px 20px rgba(108, 99, 255, 0.06);
        margin-top: 2rem;
        max-width: 700px;
        margin-left: auto;
        margin-right: auto;
    }
    .welcome-icon {
        font-size: 3.5rem;
        margin-bottom: 0.8rem;
    }
    .welcome-title {
        font-size: 1.6rem;
        font-weight: 700;
        color: #1e1e2f;
        margin-bottom: 0.5rem;
    }
    .welcome-desc {
        font-size: 0.95rem;
        color: #6b7280;
        line-height: 1.7;
        margin-bottom: 2rem;
    }
    .welcome-steps {
        display: flex;
        justify-content: center;
        gap: 1.5rem;
        flex-wrap: wrap;
    }
    .w-step {
        text-align: center;
        padding: 1rem;
        min-width: 120px;
        background: #f8f7ff;
        border-radius: 14px;
        transition: all 0.3s ease;
    }
    .w-step:hover {
        transform: translateY(-3px);
        box-shadow: 0 4px 12px rgba(108, 99, 255, 0.1);
    }
    .w-step-icon { font-size: 2rem; margin-bottom: 0.3rem; }
    .w-step-num {
        display: inline-block;
        width: 24px; height: 24px; line-height: 24px;
        border-radius: 50%;
        background: linear-gradient(135deg, #6c63ff, #8b5cf6);
        color: white; font-size: 0.75rem; font-weight: 700;
        margin-bottom: 0.3rem;
    }
    .w-step-text { font-size: 0.82rem; color: #4b5563; }

    /* ===== 刷题区域 ===== */
    .quiz-container {
        background: #ffffff;
        border-radius: 16px;
        padding: 1.5rem;
        box-shadow: 0 1px 8px rgba(0,0,0,0.04);
    }
    .quiz-meta {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        margin-bottom: 0.8rem;
        flex-wrap: wrap;
    }
    .quiz-tag {
        display: inline-block;
        padding: 3px 10px;
        border-radius: 6px;
        font-size: 0.75rem;
        font-weight: 600;
    }
    .quiz-tag.progress-tag {
        background: #ede9fe; color: #6c63ff;
    }
    .quiz-tag.concept-tag {
        background: #fce7f3; color: #db2777;
    }
    .quiz-stem-text {
        font-size: 1rem;
        font-weight: 600;
        color: #1e1e2f;
        line-height: 1.7;
        margin: 0.8rem 0 1rem 0;
        padding: 0.8rem 1rem;
        background: #f8f7ff;
        border-radius: 10px;
        border: 1px solid #ede9fe;
    }

    /* ===== 反馈卡片 ===== */
    .feedback-card {
        border-radius: 12px;
        padding: 1rem 1.2rem;
        margin: 0.8rem 0;
        line-height: 1.7;
        font-size: 0.92rem;
    }
    .feedback-card.correct {
        background: linear-gradient(135deg, #ecfdf5, #d1fae5);
        border: 1px solid #6ee7b7;
        color: #065f46;
    }
    .feedback-card.wrong {
        background: linear-gradient(135deg, #fef2f2, #fee2e2);
        border: 1px solid #fca5a5;
        color: #991b1b;
    }

    /* ===== 主区域按钮 ===== */
    .stButton > button {
        border-radius: 10px !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }
    .stButton > button:hover {
        transform: translateY(-1px) !important;
    }

    /* ===== Tab / Radio 选择器美化 ===== */
    .stRadio > div[role="radiogroup"] {
        gap: 0 !important;
        background: #f1f0fb;
        border-radius: 10px;
        padding: 3px;
    }
    .stRadio > div[role="radiogroup"] > label {
        border-radius: 8px !important;
        padding: 0.4rem 1rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }

    /* ===== 成功提示条 ===== */
    .gen-success-bar {
        background: linear-gradient(135deg, #6c63ff 0%, #8b5cf6 100%);
        color: white;
        padding: 0.8rem 1.2rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 0.92rem;
        box-shadow: 0 4px 14px rgba(108, 99, 255, 0.25);
        margin-bottom: 1rem;
    }

    /* ===== 聊天消息 ===== */
    [data-testid="stChatMessage"] {
        border-radius: 14px !important;
        box-shadow: 0 1px 4px rgba(0,0,0,0.03) !important;
    }

    /* ===== 分隔线 ===== */
    .section-divider {
        border: none;
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(108, 99, 255, 0.15), transparent);
        margin: 1.2rem 0;
    }

    /* ===== 错题循环标题 ===== */
    .remedial-header {
        background: linear-gradient(135deg, #fef3c7, #fde68a);
        border: 1px solid #fbbf24;
        border-radius: 10px;
        padding: 0.6rem 1rem;
        font-weight: 600;
        color: #92400e;
        font-size: 0.92rem;
        margin-bottom: 0.8rem;
    }
</style>
""", unsafe_allow_html=True)


# =========================================================
# 应用状态初始化
# =========================================================
def init_state():
    """初始化所有需要的状态变量"""
    
    # 课程名称
    if "course_name" not in st.session_state:
        st.session_state.course_name = "通用课程"

    # 是否正在执行生成任务
    if "run_generation" not in st.session_state:
        st.session_state.run_generation = False

    # 知识清单数据（JSON 格式）
    if "review_pack" not in st.session_state:
        st.session_state.review_pack = None

    # 知识清单展示内容（Markdown 格式）
    if "result_content" not in st.session_state:
        st.session_state.result_content = ""

    # 答疑对话历史
    if "qa_messages" not in st.session_state:
        st.session_state.qa_messages = []

    # 题组缓存（保证至少有一套题可用）
    if "question_sets" not in st.session_state:
        st.session_state.question_sets = []

    # 题组生成错误记录
    if "next_set_error" not in st.session_state:
        st.session_state.next_set_error = ""

    # 生成锁（防止重复点击"开始学习"）
    if "is_generating" not in st.session_state:
        st.session_state.is_generating = False

    # 学习模式（0基础 / 高玩）
    if "study_mode" not in st.session_state:
        st.session_state.study_mode = UI["MODE_BEGINNER"]

    # 刷题状态机
    if "quiz" not in st.session_state:
        st.session_state.quiz = {
            "active": False,              # 是否正在刷题
            "phase": "main",              # main=主卷10题 / remedial=错题循环
            "questions": [],              # 当前题组
            "idx": 0,                     # 当前题目索引
            "last_feedback": None,        # 上一题的批改结果
            "await_next": False,          # 是否等待用户点"下一题"
            "wrong_concepts": {},         # 记录每个考点的错误次数
            "concept_mastery": {},        # 记录每个考点是否已掌握
            "remedial_queue": [],         # 错题循环队列
            "current_set_id": None,       # 当前题组 ID
            "needs_prepare_next": False,  # 是否需要准备下一组题
        }


init_state()


# =========================================================
# 工具函数
# =========================================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    从 PDF 文件提取文字
    适用于可以复制文字的 PDF
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def pdf_pages_to_base64_images(file_bytes: bytes, max_pages: int = 5):
    """
    把 PDF 的前几页转成图片（Base64 格式）
    适用于扫描件或无法提取文字的 PDF
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images = []
    for i in range(min(len(doc), max_pages)):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=72)
        b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")
        images.append(f"data:image/png;base64,{b64}")
    doc.close()
    return images


def generate_word_file_from_markdown(course_name: str, md: str):
    """
    把 Markdown 格式的知识清单转成 Word 文档
    用于下载功能
    """
    doc = Document()
    doc.add_heading(f'知识清单（{course_name}）', 0)
    for line in md.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line.replace('**', '').replace('__', ''))
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def safe_extract_json(text: str):
    """
    从 AI 返回的文本中提取 JSON
    AI 有时会在 JSON 外面加一些说明文字，这个函数会自动过滤
    """
    if not text:
        return None
    # 找到第一个完整的 JSON 对象
    m = re.search(r'(\{.*\})', text, re.S)
    if not m:
        return None
    raw = m.group(1)
    try:
        return json.loads(raw)
    except:
        # 修复常见错误：去掉多余的逗号
        raw2 = re.sub(r',\s*([\]}])', r'\1', raw)
        try:
            return json.loads(raw2)
        except:
            return None


def concepts_to_markdown(pack: dict) -> str:
    """
    把结构化的知识清单转成 Markdown 格式
    用于在页面上展示
    """
    lines = [UI["MD_TITLE"]]
    if not pack or "concepts" not in pack:
        return "\n".join(lines)

    for c in pack["concepts"]:
        title = c.get("title", "未命名考点")
        explain = c.get("explain", "")
        example = c.get("example", "")
        score = c.get("final_score", c.get("model_score", 0))
        star = "⭐ " if c.get("is_star") else ""
        lines.append(f"## {star}{title}{UI['MD_SCORE'].format(score=score)}")
        if explain:
            lines.append(f"{UI['MD_EXPLAIN']}{explain}")
        if example:
            lines.append(f"{UI['MD_EXAMPLE']}{example}")
        lines.append("")
    return "\n".join(lines)


def concepts_to_html_cards(pack: dict) -> str:
    """
    把结构化的知识清单渲染为 HTML 卡片
    用于在页面上美观展示
    """
    if not pack or "concepts" not in pack:
        return ""

    html_parts = []
    for c in pack["concepts"]:
        title = c.get("title", "未命名考点")
        explain = c.get("explain", "")
        example = c.get("example", "")
        score = c.get("final_score", c.get("model_score", 0))
        is_star = c.get("is_star", False)

        # 根据分数选择徽章颜色
        if score >= 80:
            badge_cls = "hot"
        elif score >= 50:
            badge_cls = "warm"
        else:
            badge_cls = "cool"

        card_cls = "kc-card star" if is_star else "kc-card"
        star_icon = "&#11088; " if is_star else ""

        explain_html = ""
        if explain:
            explain_html = f'<div class="kc-explain"><span class="lbl">解释：</span>{explain}</div>'

        example_html = ""
        if example:
            example_html = f'<div class="kc-example">&#128161; 例子：{example}</div>'

        html_parts.append(f"""
        <div class="{card_cls}">
            <div class="kc-header">
                <div class="kc-title">{star_icon}{title}</div>
                <span class="kc-badge {badge_cls}">{score} 分</span>
            </div>
            {explain_html}
            {example_html}
        </div>
        """)

    return "\n".join(html_parts)


def is_beginner_mode(mode_str: str) -> bool:
    """判断用户选择的是 0基础模式 还是 高玩模式"""
    return mode_str.strip().startswith("俺0基础")


# =========================================================
# AI 调用函数
# =========================================================
def _set_api_key(api_key: str):
    """设置通义千问 API Key"""
    dashscope.api_key = api_key.strip().replace("：", "").replace(":", "")


def call_qwen_vl_vision(images, api_key: str) -> str:
    """
    调用视觉模型识别 PDF 图片
    用于处理扫描件或无法提取文字的 PDF
    """
    _set_api_key(api_key)
    content = [{"text": "分析PPT截图，提取核心知识点（尽量结构化输出）。"}] + [{"image": img} for img in images]
    try:
        resp = dashscope.MultiModalConversation.call(
            model='qwen-vl-max',
            messages=[{"role": "user", "content": content}]
        )
        if resp.status_code == 200:
            return resp.output.choices[0]['message']['content'][0]['text']
        return ""
    except:
        return ""


def _concepts_system_prompt(mode_str: str) -> str:
    """
    生成知识清单的系统提示词
    根据用户选择的模式（0基础/高玩）调整输出风格
    """
    beginner = is_beginner_mode(mode_str)

    # 不同模式下的知识点数量要求
    n_req = "20-30" if beginner else "12-18"
    
    # 不同模式下的解释风格
    if beginner:
        style_req = """
解释要求：
- 用大白话，像给朋友讲故事一样
- 分步骤，一步步说清楚
- 例子必须来自日常生活：比如买菜、坐地铁、刷抖音、点外卖、打游戏等
- 避免抽象概念，多用"就像..."、"好比..."这样的比喻
"""
    else:
        style_req = """
解释要求：
- 简洁有力，直击考点
- 例子可以更专业，但要实用、能记住
- 例子要能帮助快速回忆考点
"""

    return f"""
你是一个"复习包结构化生成器"。请基于用户课件与（可选）真题，输出严格 JSON（不要 markdown，不要解释）。

JSON schema（必须严格遵守）：
{{
  "course_name": "课程名称",
  "concepts": [
    {{
      "id": "c1",
      "title": "概念/考点标题（短）",
      "explain": "一句话到三句话解释（清晰、可背）",
      "example": "一个实际例子（短）",
      "is_star": true/false,
      "model_score": 0-100
    }}
  ]
}}

要求：
1) concepts 输出 {n_req} 条（宁可少而精，覆盖全课）。
2) 如果提供了真题，请将与真题高度一致/同类高频考点标记 is_star=true；若无真题，全部 is_star=false。
3) model_score 表示"你判断的可能考到程度"，0-100。
4) 风格要求：{style_req}
5) 只输出 JSON，不要任何多余字段。
""".strip()


def call_qwen_generate_concepts_json_stream(main_text: str, exam_text: str, api_key: str, mode_str: str, on_progress=None):
    """
    流式生成知识清单
    边生成边调用 on_progress 回调函数更新进度
    """
    _set_api_key(api_key)
    has_exam = bool(exam_text and exam_text.strip())

    user_content = f"""
真题：{"有" if has_exam else "无"}
【课件正文】：
{main_text[:14000]}
【真题正文】：
{exam_text[:6000] if has_exam else ""}
""".strip()

    system_prompt = _concepts_system_prompt(mode_str)

    full_text = ""
    last_len = 0
    try:
        stream = dashscope.Generation.call(
            model='qwen-plus',
            messages=[
                {'role': Role.SYSTEM, 'content': system_prompt},
                {'role': Role.USER, 'content': user_content},
            ],
            result_format='message',
            stream=True,
            incremental_output=True
        )

        for chunk in stream:
            piece = ""
            # 从不同格式的返回中提取内容
            try:
                piece = chunk.output.choices[0]['message']['content']
            except Exception:
                try:
                    piece = chunk["output"]["choices"][0]["message"]["content"]
                except Exception:
                    piece = ""

            if isinstance(piece, str) and piece:
                full_text += piece
            elif isinstance(piece, list) and piece:
                try:
                    full_text += "".join([p for p in piece if isinstance(p, str)])
                except Exception:
                    pass

            # 调用进度回调
            if on_progress and len(full_text) != last_len:
                last_len = len(full_text)
                on_progress(last_len)

        return full_text, None
    except Exception as e:
        return "", str(e)


def postprocess_concepts_json(data: dict, exam_text: str):
    """
    处理生成的知识清单 JSON
    计算最终评分并排序
    
    评分规则：
    - 有真题：60%看是否真题同类 + 40%看模型判断
    - 无真题：100%看模型判断
    """
    if not data or "concepts" not in data:
        return None, "JSON 缺少 concepts"

    has_exam = bool(exam_text and exam_text.strip())
    concepts = data["concepts"]

    for c in concepts:
        is_star = bool(c.get("is_star")) if has_exam else False
        ms = c.get("model_score", 0)
        try:
            ms = int(ms)
        except:
            ms = 0
        ms = max(0, min(100, ms))

        if has_exam:
            final = round(0.6 * (100 if is_star else 0) + 0.4 * ms)
        else:
            final = ms
            c["is_star"] = False

        c["model_score"] = ms
        c["final_score"] = int(final)

    data["course_name"] = data.get("course_name") or "通用课程"
    # 按重点指数从高到低排序
    data["concepts"] = sorted(concepts, key=lambda x: x.get("final_score", 0), reverse=True)
    return data, None


def call_qwen_generate_question_set(review_pack: dict, api_key: str, mode_str: str):
    """
    生成一组 10 道题
    题型：7单选 + 2简答 + 1论述
    按知识点重要度从高到低出题
    """
    _set_api_key(api_key)
    concepts = review_pack.get("concepts", []) if review_pack else []

    # 取前 12 个高频知识点作为出题范围
    top = concepts[:12]
    top_payload = [
        {
            "id": c.get("id"),
            "title": c.get("title"),
            "explain": c.get("explain"),
            "example": c.get("example"),
            "is_star": c.get("is_star", False),
            "final_score": c.get("final_score", 0),
        } for c in top
    ]

    beginner = is_beginner_mode(mode_str)
    
    # 不同模式下的题目难度要求
    difficulty_req = (
        "难度偏基础：以定义、辨析、理解为主，干扰项不要太阴间。"
        if beginner else
        "难度偏高：增加易混淆选项与边界条件，考察概念辨析与应用陷阱。"
    )
    style_req = (
        "题干更清晰直接，避免过长；简答/论述给出明确作答方向。"
        if beginner else
        "题干更像真题，设置反直觉陷阱但不允许超纲；简答/论述更看结构与推理。"
    )

    system_prompt = f"""
你是一个"高频考点出题器"。请严格输出 JSON（不要 markdown，不要解释）。
目标：生成一组 10 道题，按"最可能考"到"较不可能考"排序。

规则：
1) 必须严格生成：7道单选题 + 2道简答题 + 1道综合论述题。
2) 每道题必须绑定一个概念 concept_id，并尽量覆盖不同概念。
3) 单选题提供 A/B/C/D 四个选项；answer 必须是 "A"|"B"|"C"|"D"。
4) 简答/论述题 answer 用"要点列表"的形式给评分要点。
5) 解析 explanation 要简洁，允许少量扩展解释，但考点必须来自给定概念。
6) questions 必须按概念 final_score 高->低排序优先绑定。
7) {difficulty_req}
8) {style_req}

JSON schema：
{{
  "set_id": "string",
  "questions": [
    {{
      "qid": 1,
      "type": "single"|"short"|"essay",
      "concept_id": "c1",
      "concept_title": "概念标题",
      "stem": "题干",
      "options": {{"A":"...", "B":"...", "C":"...", "D":"..."}},
      "answer": "A" 或 ["要点1","要点2"],
      "explanation": "解析（用于答错时）"
    }}
  ]
}}
""".strip()

    user_content = json.dumps({"top_concepts": top_payload}, ensure_ascii=False)

    try:
        resp = dashscope.Generation.call(
            model='qwen-plus',
            messages=[
                {'role': Role.SYSTEM, 'content': system_prompt},
                {'role': Role.USER, 'content': user_content},
            ],
            result_format='message'
        )
        if resp.status_code != 200:
            return None, f"Error: {resp.message}"

        raw = resp.output.choices[0]['message']['content']
        data = safe_extract_json(raw)
        if not data or "questions" not in data:
            return None, "Error: 题组 JSON 解析失败"

        if not data.get("set_id"):
            data["set_id"] = f"set_{int(time.time())}_{random.randint(1000,9999)}"

        return data, None
    except Exception as e:
        return None, str(e)


def call_qwen_grade_answer(question: dict, user_answer: str, review_pack: dict, api_key: str):
    """
    批改用户答案
    策略：
    - 答对：告诉用户为什么做对了
    - 答错：给出正确答案和解析
    """
    _set_api_key(api_key)
    concept_id = question.get("concept_id")
    concept_ctx = None
    for c in (review_pack.get("concepts") or []):
        if c.get("id") == concept_id:
            concept_ctx = c
            break
    concept_ctx = concept_ctx or {}

    system_prompt = """
你是一个"严格锚定知识清单的阅卷老师"。请只输出 JSON（不要 markdown，不要解释）。
要求：
1) 判断用户答案正确与否。
2) 若正确：只输出 why_correct（告诉用户他为什么做对了）。
3) 若错误：输出 correct_answer + explanation。
4) 允许少量扩展解释，但考点必须来自 concept_context。
JSON schema：
{
  "is_correct": true/false,
  "why_correct": "string",
  "correct_answer": "string",
  "explanation": "string"
}
注意：正确时 correct_answer/explanation 允许为空；错误时 why_correct 允许为空。
""".strip()

    payload = {
        "question": question,
        "user_answer": user_answer,
        "concept_context": {
            "title": concept_ctx.get("title", ""),
            "explain": concept_ctx.get("explain", ""),
            "example": concept_ctx.get("example", ""),
            "is_star": concept_ctx.get("is_star", False),
        }
    }

    try:
        resp = dashscope.Generation.call(
            model='qwen-plus',
            messages=[
                {'role': Role.SYSTEM, 'content': system_prompt},
                {'role': Role.USER, 'content': json.dumps(payload, ensure_ascii=False)},
            ],
            result_format='message'
        )
        if resp.status_code != 200:
            return None, f"Error: {resp.message}"

        raw = resp.output.choices[0]['message']['content']
        data = safe_extract_json(raw)
        if not data or "is_correct" not in data:
            return None, "Error: 批改 JSON 解析失败"
        return data, None
    except Exception as e:
        return None, str(e)


def call_qwen_generate_similar_question(concept_id: str, preferred_type: str, review_pack: dict, api_key: str):
    """
    生成同考点的类似题
    用于错题循环：同一个知识点换个问法再考一遍
    """
    _set_api_key(api_key)
    concept_ctx = None
    for c in (review_pack.get("concepts") or []):
        if c.get("id") == concept_id:
            concept_ctx = c
            break
    concept_ctx = concept_ctx or {}

    qtype = preferred_type if preferred_type in ["single", "short", "essay"] else "single"

    system_prompt = """
你是一个"错题再训练出题器"。请严格输出 JSON（不要 markdown，不要解释）。
目标：出一道"同考点、不同问法"的类似题，用于纠错训练。
JSON schema：
{
  "qid": 999,
  "type": "single"|"short"|"essay",
  "concept_id": "c1",
  "concept_title": "概念标题",
  "stem": "题干",
  "options": {"A":"...", "B":"...", "C":"...", "D":"..."},
  "answer": "A" 或 ["要点1","要点2"],
  "explanation": "解析"
}
""".strip()

    payload = {
        "type": qtype,
        "concept_context": {
            "id": concept_id,
            "title": concept_ctx.get("title", ""),
            "explain": concept_ctx.get("explain", ""),
            "example": concept_ctx.get("example", ""),
        }
    }

    try:
        resp = dashscope.Generation.call(
            model='qwen-plus',
            messages=[
                {'role': Role.SYSTEM, 'content': system_prompt},
                {'role': Role.USER, 'content': json.dumps(payload, ensure_ascii=False)},
            ],
            result_format='message'
        )
        if resp.status_code != 200:
            return None, f"Error: {resp.message}"

        raw = resp.output.choices[0]['message']['content']
        data = safe_extract_json(raw)
        if not data or "stem" not in data:
            return None, "Error: 类似题 JSON 解析失败"

        data["concept_id"] = concept_id
        data["concept_title"] = concept_ctx.get("title", data.get("concept_title", ""))
        return data, None
    except Exception as e:
        return None, str(e)


def call_qwen_qa(user_question: str, review_pack: dict, api_key: str):
    """
    答疑功能
    只能基于知识清单回答，不会引入新的考点
    """
    _set_api_key(api_key)
    concepts = (review_pack.get("concepts") or [])[:25]
    ctx = [{"title": c.get("title", ""), "explain": c.get("explain", ""), "example": c.get("example", "")} for c in concepts]

    system_prompt = """
你是一个"知识清单内答疑助教"。请基于给定知识清单内容回答问题。
要求：
1) 优先引用知识清单的解释与例子进行回答。
2) 允许少量扩展解释，但不要引入与知识清单无关的新考点。
3) 如果用户问到清单未覆盖内容，说明"知识清单未包含"，并建议补充课件或生成更完整清单。
""".strip()

    payload = {"concepts": ctx, "question": user_question}
    try:
        resp = dashscope.Generation.call(
            model='qwen-plus',
            messages=[
                {'role': Role.SYSTEM, 'content': system_prompt},
                {'role': Role.USER, 'content': json.dumps(payload, ensure_ascii=False)},
            ],
            result_format='message'
        )
        if resp.status_code == 200:
            return resp.output.choices[0]['message']['content']
        return "Error"
    except:
        return "Error"


# =========================================================
# 刷题状态机
# =========================================================
def quiz_reset():
    """重置刷题状态"""
    st.session_state.quiz = {
        "active": False,
        "phase": "main",
        "questions": [],
        "idx": 0,
        "last_feedback": None,
        "await_next": False,
        "wrong_concepts": {},
        "concept_mastery": {},
        "remedial_queue": [],
        "current_set_id": None,
        "needs_prepare_next": False,
    }


def quiz_start_with_set(question_set: dict):
    """开始一组新题"""
    quiz_reset()
    st.session_state.quiz["active"] = True
    st.session_state.quiz["phase"] = "main"
    st.session_state.quiz["questions"] = question_set.get("questions", [])
    st.session_state.quiz["idx"] = 0
    st.session_state.quiz["current_set_id"] = question_set.get("set_id")


def quiz_current_question():
    """获取当前正在做的题"""
    qz = st.session_state.quiz
    if not qz["active"]:
        return None
    if qz["idx"] < 0 or qz["idx"] >= len(qz["questions"]):
        return None
    return qz["questions"][qz["idx"]]


def register_wrong_concept(question: dict):
    """记录答错的考点"""
    cid = question.get("concept_id", "unknown")
    qtype = question.get("type", "single")
    wrong = st.session_state.quiz["wrong_concepts"].get(cid)
    if not wrong:
        st.session_state.quiz["wrong_concepts"][cid] = {"attempts": 1, "stuck": False, "preferred_type": qtype}
    else:
        wrong["attempts"] += 1
    st.session_state.quiz["concept_mastery"][cid] = False


def register_correct_concept(question: dict):
    """记录答对的考点"""
    cid = question.get("concept_id", "unknown")
    st.session_state.quiz["concept_mastery"][cid] = True


def build_remedial_queue():
    """构建错题循环队列"""
    wrong = st.session_state.quiz["wrong_concepts"]
    queue = []
    for cid, meta in wrong.items():
        if meta.get("stuck"):
            continue
        if st.session_state.quiz["concept_mastery"].get(cid) is False:
            queue.append(cid)
    st.session_state.quiz["remedial_queue"] = queue


def move_to_next_question_or_phase():
    """移动到下一题或下一阶段"""
    qz = st.session_state.quiz
    qz["last_feedback"] = None
    qz["await_next"] = False

    if qz["phase"] == "main":
        qz["idx"] += 1
        if qz["idx"] >= len(qz["questions"]):
            # 主卷做完了，准备进入错题循环或结束
            # 这时候偷偷生成下一组题（无论有没有错题）
            qz["needs_prepare_next"] = True  # 标记需要准备下一组题
            
            build_remedial_queue()
            if qz["remedial_queue"]:
                qz["phase"] = "remedial"
                qz["idx"] = 0
                qz["questions"] = []
            else:
                qz["active"] = False
    else:
        build_remedial_queue()
        if not qz["remedial_queue"]:
            qz["active"] = False


def get_concept_snippet(concept_id: str, review_pack: dict) -> str:
    """获取某个考点的简短说明"""
    for c in (review_pack.get("concepts") or []):
        if c.get("id") == concept_id:
            return f"**{c.get('title','')}**：{c.get('explain','')}"
    return ""


# =========================================================
# 侧边栏界面
# =========================================================
with st.sidebar:
    st.title(UI["SIDEBAR_TITLE"])

    # API Key 配置
    api_key = st.secrets.get("DASHSCOPE_API_KEY", os.getenv("DASHSCOPE_API_KEY", ""))
    if api_key:
        st.success(UI["DEV_MODE"])
    else:
        api_key = st.text_input("API Key", type="password")

    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

    # 工具箱（只在生成知识清单后显示）
    if st.session_state.review_pack:
        st.markdown(UI["TOOLBOX"])

        # 考考我按钮
        quiz_btn_clicked = st.button(UI["QUIZ_BTN"], use_container_width=True)
        
        if quiz_btn_clicked:
            if not api_key:
                st.error("ERROR_CODE: SYS_001 - 系统配置缺失")
            elif not st.session_state.question_sets:
                # 有错误信息：显示错误码
                if st.session_state.next_set_error:
                    st.error(f"ERROR_CODE: {st.session_state.next_set_error}")
                else:
                    # 不太可能，但保险起见
                    st.warning("😅 题目还在准备中，请稍等片刻再试")
            else:
                # 有题了，开始刷题
                qs = st.session_state.question_sets.pop(0)
                quiz_start_with_set(qs)
                st.rerun()

        # 下载按钮
        docx_file = generate_word_file_from_markdown(
            st.session_state.course_name,
            st.session_state.result_content
        )
        st.download_button(
            label=UI["DOWNLOAD_BTN"],
            data=docx_file,
            file_name=f"{st.session_state.course_name}_知识清单.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

    # 设置区域
    st.markdown(UI["SETTINGS"])
    mode = st.radio(UI["GOAL"], (UI["MODE_BEGINNER"], UI["MODE_PRO"]), index=0)
    st.session_state.study_mode = mode

    st.caption(UI["UPLOAD_TIP"])
    uploaded_files = st.file_uploader(UI["UPLOAD_MAIN"], type=["pdf"], accept_multiple_files=True)

    st.caption(UI["UPLOAD_TIP"])
    uploaded_exams = st.file_uploader(UI["UPLOAD_EXAM"], type=["pdf"], key="exam")

    # 开始学习按钮
    start_clicked = st.button(
        UI["START_BTN"],
        type="primary",
        use_container_width=True,
        disabled=st.session_state.is_generating or (st.session_state.review_pack is not None),
        key="start_learning_btn"
    )


# =========================================================
# 主界面
# =========================================================
st.markdown(f'<div class="main-hero-title">{UI["MAIN_TITLE"]}</div>', unsafe_allow_html=True)
st.markdown('<div class="main-hero-sub">上传课件，AI 自动提取核心考点 / 生成高频题库 / 智能答疑</div>', unsafe_allow_html=True)

# 进度显示占位符
progress_container = st.empty()
status_container = st.empty()

# =========================================================
# 核心生成流程
# =========================================================
if start_clicked:
    st.session_state.is_generating = True
    st.session_state.run_generation = True
    st.rerun()

if st.session_state.run_generation:
    try:
        # 检查是否上传了文件
        if not uploaded_files:
            st.session_state.run_generation = False
            st.session_state.is_generating = False
            st.error("请先上传课件 PDF")
            st.stop()

        # 显示初始进度
        progress_bar = progress_container.progress(0)
        status_text = status_container.text(UI["GEN_PROGRESS"])

        # 提取课件文字
        main_text = ""
        for uploaded_file in uploaded_files:
            file_bytes = uploaded_file.read()
            text = extract_text_from_pdf(file_bytes)
            
            # 如果文字太少，尝试用视觉模型识别
            if len(text.strip()) < 100:
                images = pdf_pages_to_base64_images(file_bytes, max_pages=5)
                if images and api_key:
                    ocr_text = call_qwen_vl_vision(images, api_key)
                    text += "\n" + ocr_text
            
            main_text += "\n" + text

        if len(main_text.strip()) < 50:
            st.session_state.run_generation = False
            st.session_state.is_generating = False
            progress_container.empty()
            status_container.empty()
            st.error("ERROR_CODE: PDF_EXT_001 - 未能从 PDF 中提取到足够内容")
            st.stop()

        # 提取真题文字（如果有）
        exam_text = ""
        if uploaded_exams:
            exam_bytes = uploaded_exams.read()
            exam_text = extract_text_from_pdf(exam_bytes)
            
            # 真题也尝试 OCR（如果文字太少）
            if len(exam_text.strip()) < 100:
                exam_images = pdf_pages_to_base64_images(exam_bytes, max_pages=5)
                if exam_images and api_key:
                    exam_ocr = call_qwen_vl_vision(exam_images, api_key)
                    exam_text += "\n" + exam_ocr

        # ========================================
        # 阶段1：流式生成知识清单 (0-70%)
        # ========================================
        target_length = 5000  # 预估目标长度
        
        def update_progress(current_len):
            """更新进度条的回调函数（只到70%）"""
            percentage = min(int((current_len / target_length) * 70), 70)
            progress_bar.progress(percentage)
            
            # 计算预计剩余时间
            if percentage < 20:
                time_left = 2
            elif percentage < 50:
                time_left = 1
            else:
                time_left = 0.5
            
            status_text.text(f"{UI['GEN_PROGRESS']} {percentage}% | {UI['GEN_ESTIMATE'].format(time=int(time_left))}")

        # 调用 AI 生成知识清单
        raw_json, err = call_qwen_generate_concepts_json_stream(
            main_text=main_text,
            exam_text=exam_text,
            api_key=api_key,
            mode_str=st.session_state.study_mode,
            on_progress=update_progress
        )

        if err:
            st.session_state.run_generation = False
            st.session_state.is_generating = False
            progress_container.empty()
            status_container.empty()
            st.error(f"ERROR_CODE: KC_GEN_001 - {err}")
            st.stop()

        # 解析 JSON
        data = safe_extract_json(raw_json)
        if not data:
            st.session_state.run_generation = False
            st.session_state.is_generating = False
            progress_container.empty()
            status_container.empty()
            st.error(f"ERROR_CODE: KC_GEN_002 - 无法解析 AI 返回的 JSON")
            st.stop()

        # 后处理（计算评分、排序）
        pack, post_err = postprocess_concepts_json(data, exam_text)
        if post_err:
            st.session_state.run_generation = False
            st.session_state.is_generating = False
            progress_container.empty()
            status_container.empty()
            st.error(f"ERROR_CODE: KC_GEN_003 - {post_err}")
            st.stop()

        # 保存知识清单到状态
        st.session_state.review_pack = pack
        st.session_state.course_name = pack.get("course_name", "通用课程")
        st.session_state.result_content = concepts_to_markdown(pack)

        # ========================================
        # 阶段2：偷偷生成第一组题 (70-100%)
        # ========================================
        # 记录开始时间，用于动态调整进度条速度
        import time as time_module
        stage2_start = time_module.time()
        
        # 启动题组生成（异步，但我们会等它完成）
        mode_str = st.session_state.study_mode
        first_set, set_err = call_qwen_generate_question_set(pack, api_key, mode_str)
        
        stage2_duration = time_module.time() - stage2_start
        
        # 根据实际耗时调整进度条速度
        if stage2_duration < 3:
            # 太快了，强制延迟到 3 秒让进度条看起来真实
            remaining_time = 3 - stage2_duration
            steps = 30  # 从 70% 到 100%
            delay_per_step = remaining_time / steps
        else:
            # 已经够慢了，快速走完进度条
            delay_per_step = 0.05
        
        # 平滑推进进度条 70% → 100%
        for i in range(71, 101):
            progress_bar.progress(i)
            time_left = max(0, int((100 - i) * delay_per_step / 60))  # 转成分钟
            status_text.text(f"{UI['GEN_PROGRESS']} {i}% | {UI['GEN_ESTIMATE'].format(time=time_left if time_left > 0 else 0.5)}")
            time.sleep(delay_per_step)
        
        # 处理题组生成结果
        if set_err:
            st.session_state.next_set_error = f"QS_GEN_001 - {set_err}"
        elif not first_set:
            st.session_state.next_set_error = "QS_GEN_002 - 生成的题组为空"
        else:
            st.session_state.question_sets.append(first_set)
            st.session_state.next_set_error = ""  # 清空错误
        
        # ========================================
        # 完成动画
        # ========================================
        progress_bar.progress(100)
        status_text.text("")
        
        # 成功提示（带动画效果）
        success_msg = status_container.success(UI["GEN_SUCCESS_TEMP"])
        time.sleep(1.5)  # 显示 1.5 秒
        
        # 清理进度显示
        progress_container.empty()
        status_container.empty()

        # 重置生成状态
        st.session_state.run_generation = False
        st.session_state.is_generating = False
        
        # 刷新页面显示结果
        st.rerun()

    except Exception as e:
        st.session_state.run_generation = False
        st.session_state.is_generating = False
        progress_container.empty()
        status_container.empty()
        st.error(f"{UI['GEN_FAIL']}{str(e)}")
        st.stop()


# =========================================================
# 内容展示区域
# =========================================================
if st.session_state.review_pack:
    # 成功提示 - 使用自定义样式
    n_concepts = len(st.session_state.review_pack.get("concepts", []))
    st.markdown(
        f'<div class="gen-success-bar">&#10004;&#65039; 《{st.session_state.course_name}》知识清单已生成'
        f'（共 {n_concepts} 个核心考点）&nbsp;&mdash;&nbsp;侧边栏可下载 / 考考我</div>',
        unsafe_allow_html=True
    )

    # 分两栏显示
    col1, col2 = st.columns([6, 4])

    # 左栏：知识清单（HTML卡片）
    with col1:
        st.subheader(UI["LEFT_HEADER"])
        cards_html = concepts_to_html_cards(st.session_state.review_pack)
        st.markdown(cards_html, unsafe_allow_html=True)

    # 右栏：答疑 / 刷题
    with col2:
        st.subheader(UI["RIGHT_HEADER"])
        tab = st.radio(UI["MODE_SWITCH"], (UI["MODE_QA"], UI["MODE_QUIZ"]), horizontal=True)

        # -------------------------
        # 答疑模式
        # -------------------------
        if tab == UI["MODE_QA"]:
            # 显示对话历史
            for msg in st.session_state.qa_messages:
                with st.chat_message(msg["role"]):
                    st.write(msg["content"])

            # 输入框
            q = st.chat_input(UI["QA_INPUT"])
            if q:
                st.session_state.qa_messages.append({"role": "user", "content": q})
                with st.chat_message("user"):
                    st.write(q)

                # 调用 AI 回答
                with st.spinner(UI["QA_SPINNER"]):
                    a = call_qwen_qa(q, st.session_state.review_pack, api_key)

                st.session_state.qa_messages.append({"role": "assistant", "content": a})
                with st.chat_message("assistant"):
                    st.write(a)

        # -------------------------
        # 刷题模式
        # -------------------------
        else:
            qz = st.session_state.quiz

            # 检查是否需要生成下一组题（主卷做完后触发）
            if qz.get("needs_prepare_next", False):
                qz["needs_prepare_next"] = False  # 清除标记
                
                # 显示"正在整理错题"掩护生成过程
                with st.spinner("正在整理错题，马上开始..."):
                    # 偷偷生成下一组题
                    pack = st.session_state.review_pack
                    mode_str = st.session_state.study_mode
                    
                    next_set, next_err = call_qwen_generate_question_set(pack, api_key, mode_str)
                    
                    if next_err:
                        # 静默记录错误，不影响当前错题循环
                        st.session_state.next_set_error = f"QS_GEN_003 - {next_err}"
                    elif not next_set:
                        st.session_state.next_set_error = "QS_GEN_004 - 生成的题组为空"
                    else:
                        st.session_state.question_sets.append(next_set)
                        st.session_state.next_set_error = ""
                    
                    # 确保至少显示 3 秒，让用户感觉在认真整理
                    time.sleep(max(0, 3))
                
                # 整理完毕，继续刷题
                st.rerun()

            if not qz["active"]:
                st.info(UI["QUIZ_HINT"])
            else:
                # 错题循环阶段
                if qz["phase"] == "remedial":
                    cid = qz["remedial_queue"][0] if qz["remedial_queue"] else None
                    if not cid:
                        st.success(UI["REMEDIAL_DONE"])
                        qz["active"] = False
                        st.rerun()

                    meta = qz["wrong_concepts"].get(cid, {"attempts": 0, "preferred_type": "single", "stuck": False})
                    attempts = meta.get("attempts", 0)

                    # 错误次数达到上限
                    if attempts >= 4:
                        meta["stuck"] = True
                        qz["wrong_concepts"][cid] = meta
                        snippet = get_concept_snippet(cid, st.session_state.review_pack)
                        st.warning(UI["REMEDIAL_STUCK"])
                        if snippet:
                            st.markdown(f"{UI['REMEDIAL_REVIEW']} {snippet}")

                        qz["remedial_queue"] = [x for x in qz["remedial_queue"] if x != cid]
                        build_remedial_queue()
                        if not qz["remedial_queue"]:
                            st.success(UI["REMEDIAL_END"])
                            qz["active"] = False
                            st.rerun()
                        else:
                            st.rerun()

                    # 生成类似题
                    if "remedial_current_q" not in st.session_state or st.session_state.get("remedial_current_cid") != cid:
                        simq, serr = call_qwen_generate_similar_question(
                            concept_id=cid,
                            preferred_type=meta.get("preferred_type", "single"),
                            review_pack=st.session_state.review_pack,
                            api_key=api_key
                        )
                        if serr or not simq:
                            meta["stuck"] = True
                            qz["wrong_concepts"][cid] = meta
                            qz["remedial_queue"] = [x for x in qz["remedial_queue"] if x != cid]
                            st.rerun()

                        st.session_state.remedial_current_q = simq
                        st.session_state.remedial_current_cid = cid

                    current_q = st.session_state.remedial_current_q
                    st.markdown(
                        f'<div class="remedial-header">&#128260; 错题循环 &mdash; 直到掌握</div>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f'<div class="quiz-meta">'
                        f'<span class="quiz-tag concept-tag">{current_q.get("concept_title", "")}</span>'
                        f'<span class="quiz-tag progress-tag">已错 {meta.get("attempts", 0)}/4 次</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                # 主卷阶段
                else:
                    current_q = quiz_current_question()
                    cur_idx = qz["idx"] + 1
                    total_q = len(qz["questions"])
                    # 进度点
                    dots_html = ""
                    for di in range(total_q):
                        if di < cur_idx - 1:
                            dots_html += '<span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#34d399;margin:0 2px;"></span>'
                        elif di == cur_idx - 1:
                            dots_html += '<span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:linear-gradient(135deg,#6c63ff,#8b5cf6);margin:0 2px;"></span>'
                        else:
                            dots_html += '<span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#e5e7eb;margin:0 2px;"></span>'

                    st.markdown(f"### &#128221; 高频卷（{total_q}题）")
                    st.markdown(
                        f'<div class="quiz-meta">'
                        f'<span class="quiz-tag progress-tag">第 {cur_idx}/{total_q} 题</span>'
                        f'<span class="quiz-tag concept-tag">{current_q.get("concept_title", "")}</span>'
                        f'</div>'
                        f'<div style="margin-bottom:0.8rem;">{dots_html}</div>',
                        unsafe_allow_html=True
                    )

                # 显示题目 - 使用卡片样式
                st.markdown(
                    f'<div class="quiz-stem-text">{current_q.get("stem", "")}</div>',
                    unsafe_allow_html=True
                )

                # 显示上一题的批改结果 - 使用反馈卡片
                if qz["last_feedback"]:
                    feedback_text = qz["last_feedback"]
                    feedback_html = feedback_text.replace("\n\n", "<br>").replace("\n", "<br>")
                    if feedback_text.startswith(UI["CORRECT"]):
                        st.markdown(
                            f'<div class="feedback-card correct">{feedback_html}</div>',
                            unsafe_allow_html=True
                        )
                    elif feedback_text.startswith(UI["WRONG"]):
                        st.markdown(
                            f'<div class="feedback-card wrong">{feedback_html}</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        st.markdown(feedback_text)

                # 如果等待下一题，显示按钮
                if qz["await_next"]:
                    if st.button(UI["NEXT_BTN"], use_container_width=True):
                        if qz["phase"] == "remedial":
                            if "remedial_current_q" in st.session_state:
                                del st.session_state.remedial_current_q
                            if "remedial_current_cid" in st.session_state:
                                del st.session_state.remedial_current_cid
                        move_to_next_question_or_phase()
                        st.rerun()
                    st.stop()

                # 批改答案的统一函数
                def submit_answer(answer_text: str):
                    grade, gerr = call_qwen_grade_answer(
                        question=current_q,
                        user_answer=answer_text,
                        review_pack=st.session_state.review_pack,
                        api_key=api_key
                    )
                    if gerr or not grade:
                        qz["last_feedback"] = f"{UI['GRADE_FAIL']}{gerr}"
                        qz["await_next"] = True
                        return

                    is_correct = bool(grade.get("is_correct"))
                    if is_correct:
                        why = (grade.get("why_correct") or "").strip()
                        qz["last_feedback"] = f"{UI['CORRECT']}\n\n" + (why if why else "你抓住了关键考点。")
                        register_correct_concept(current_q)

                        if qz["phase"] == "remedial":
                            cid2 = current_q.get("concept_id", "unknown")
                            qz["concept_mastery"][cid2] = True
                            qz["remedial_queue"] = [x for x in qz["remedial_queue"] if x != cid2]
                    else:
                        ca = (grade.get("correct_answer") or "").strip()
                        ex = (grade.get("explanation") or "").strip()
                        block = f"{UI['WRONG']}\n\n"
                        if ca:
                            block += f"{UI['RIGHT_ANSWER']} {ca}\n\n"
                        block += ex if ex else UI["SUGGEST_REVIEW"]
                        qz["last_feedback"] = block
                        register_wrong_concept(current_q)

                    qz["await_next"] = True

                # 根据题型显示不同的答题界面
                qtype = current_q.get("type", "single")
                
                # 单选题：ABCD 按钮
                if qtype == "single":
                    opts = current_q.get("options", {}) or {}
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button(f"A. {opts.get('A','')}", use_container_width=True):
                            submit_answer("A")
                            st.rerun()
                        if st.button(f"C. {opts.get('C','')}", use_container_width=True):
                            submit_answer("C")
                            st.rerun()
                    with c2:
                        if st.button(f"B. {opts.get('B','')}", use_container_width=True):
                            submit_answer("B")
                            st.rerun()
                        if st.button(f"D. {opts.get('D','')}", use_container_width=True):
                            submit_answer("D")
                            st.rerun()
                
                # 简答/论述题：文本框
                else:
                    user_text = st.text_area(
                        UI["TEXTAREA_LABEL"],
                        height=120,
                        placeholder=UI["TEXTAREA_PLACEHOLDER"]
                    )
                    if st.button(UI["SUBMIT_BTN"], use_container_width=True):
                        if not user_text.strip():
                            st.warning(UI["NEED_INPUT"])
                        else:
                            submit_answer(user_text.strip())
                            st.rerun()

else:
    # 未生成知识清单时显示欢迎页
    if not st.session_state.is_generating and not st.session_state.run_generation:
        st.markdown("""
        <div class="welcome-container">
            <div class="welcome-icon">&#127819;</div>
            <div class="welcome-title">把课件扔进来，知识榨出来</div>
            <div class="welcome-desc">
                上传 PDF 课件，AI 帮你提炼核心考点、生成练习题、智能答疑<br>
                期末复习不再迷茫，高效拿分
            </div>
            <div class="welcome-steps">
                <div class="w-step">
                    <div class="w-step-icon">&#128196;</div>
                    <div class="w-step-num">1</div>
                    <div class="w-step-text">上传课件 PDF</div>
                </div>
                <div class="w-step">
                    <div class="w-step-icon">&#9889;</div>
                    <div class="w-step-num">2</div>
                    <div class="w-step-text">AI 自动提取考点</div>
                </div>
                <div class="w-step">
                    <div class="w-step-icon">&#128221;</div>
                    <div class="w-step-num">3</div>
                    <div class="w-step-text">刷题 + 答疑</div>
                </div>
                <div class="w-step">
                    <div class="w-step-icon">&#127942;</div>
                    <div class="w-step-num">4</div>
                    <div class="w-step-text">轻松应对考试</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)